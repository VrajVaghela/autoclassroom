"""
Word and PDF report rendering from one shared block list.

A document is a list of blocks the LLM produced:
  {"type": "heading",   "text": "...", "level": 1}
  {"type": "paragraph", "text": "..."}
  {"type": "code",      "text": "...", "language": "python"}
  {"type": "screenshot","command": "...", "output": "...", "caption": "..."}
  {"type": "table",     "headers": [...], "rows": [[...], ...], "caption": "..."}
  {"type": "list",      "items": [...], "ordered": false}
  {"type": "pagebreak"}

Screenshot blocks are rendered to PNG by artifacts.screenshot and embedded, so
a lab report arrives with output images already in place.
"""

import os

from .screenshot import render_terminal

CODE_FONT = "Consolas"


def normalize_blocks(blocks):
    """Drop anything malformed so one bad block can't fail the whole document."""
    clean = []
    for b in blocks or []:
        if not isinstance(b, dict):
            continue
        kind = (b.get("type") or "paragraph").lower()
        if kind in ("heading", "paragraph", "code", "screenshot", "table", "list", "pagebreak"):
            clean.append(dict(b, type=kind))
    return clean


def _screenshot_png(block, assets_dir, index):
    os.makedirs(assets_dir, exist_ok=True)
    path = os.path.join(assets_dir, f"output_{index:02d}.png")
    return render_terminal(
        path,
        block.get("command", ""),
        block.get("output", ""),
        block.get("title") or "Terminal",
    )


def write_docx(path, title, blocks, assets_dir):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    if title:
        doc.add_heading(title, level=0)

    shot_index = 0
    for block in normalize_blocks(blocks):
        kind = block["type"]

        if kind == "heading":
            level = max(1, min(4, int(block.get("level") or 1)))
            doc.add_heading(block.get("text", ""), level=level)

        elif kind == "paragraph":
            doc.add_paragraph(block.get("text", ""))

        elif kind == "list":
            style = "List Number" if block.get("ordered") else "List Bullet"
            for item in block.get("items") or []:
                doc.add_paragraph(str(item), style=style)

        elif kind == "code":
            para = doc.add_paragraph()
            run = para.add_run(block.get("text", ""))
            run.font.name = CODE_FONT
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            para.paragraph_format.left_indent = Inches(0.25)
            para.paragraph_format.space_after = Pt(10)

        elif kind == "screenshot":
            shot_index += 1
            png = _screenshot_png(block, assets_dir, shot_index)
            doc.add_picture(png, width=Inches(6.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption = block.get("caption") or f"Output {shot_index}"
            cap = doc.add_paragraph(f"Figure {shot_index}: {caption}")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].font.size = Pt(9)
            cap.runs[0].italic = True

        elif kind == "table":
            headers = block.get("headers") or []
            rows = block.get("rows") or []
            if not headers and not rows:
                continue
            ncols = len(headers) or max(len(r) for r in rows)
            table = doc.add_table(rows=0, cols=ncols)
            table.style = "Table Grid"
            if headers:
                cells = table.add_row().cells
                for i, head in enumerate(headers[:ncols]):
                    cells[i].text = str(head)
                    for para in cells[i].paragraphs:
                        for run in para.runs:
                            run.bold = True
            for row in rows:
                cells = table.add_row().cells
                for i, value in enumerate(list(row)[:ncols]):
                    cells[i].text = "" if value is None else str(value)
            if block.get("caption"):
                doc.add_paragraph(block["caption"]).runs[0].italic = True

        elif kind == "pagebreak":
            doc.add_page_break()

    os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
    doc.save(path)
    return path


def write_pdf(path, title, blocks, assets_dir):
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.lib.utils import ImageReader
    from reportlab.platypus import (Image as RLImage, ListFlowable, ListItem,
                                    PageBreak, Paragraph, Preformatted,
                                    SimpleDocTemplate, Spacer, Table, TableStyle)

    styles = getSampleStyleSheet()
    code_style = ParagraphStyle(
        "AutoCode", parent=styles["Code"], fontName="Courier", fontSize=8.5,
        leading=11, leftIndent=12, backColor=colors.HexColor("#F4F4F4"),
        borderPadding=6, spaceAfter=10,
    )
    caption_style = ParagraphStyle(
        "AutoCaption", parent=styles["Normal"], fontSize=8.5,
        alignment=TA_CENTER, textColor=colors.HexColor("#555555"), spaceAfter=12,
    )

    def esc(text):
        return (str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))

    story = []
    if title:
        story += [Paragraph(esc(title), styles["Title"]), Spacer(1, 10)]

    shot_index = 0
    for block in normalize_blocks(blocks):
        kind = block["type"]

        if kind == "heading":
            level = max(1, min(4, int(block.get("level") or 1)))
            story.append(Paragraph(esc(block.get("text", "")), styles[f"Heading{level}"]))

        elif kind == "paragraph":
            story += [Paragraph(esc(block.get("text", "")), styles["BodyText"]), Spacer(1, 4)]

        elif kind == "list":
            items = [ListItem(Paragraph(esc(i), styles["BodyText"]))
                     for i in (block.get("items") or [])]
            if items:
                story += [ListFlowable(items,
                                       bulletType="1" if block.get("ordered") else "bullet"),
                          Spacer(1, 8)]

        elif kind == "code":
            story.append(Preformatted(block.get("text", ""), code_style))

        elif kind == "screenshot":
            shot_index += 1
            png = _screenshot_png(block, assets_dir, shot_index)
            iw, ih = ImageReader(png).getSize()
            width = min(6.2 * inch, iw)
            story.append(RLImage(png, width=width, height=width * ih / iw))
            caption = block.get("caption") or f"Output {shot_index}"
            story.append(Paragraph(f"Figure {shot_index}: {esc(caption)}", caption_style))

        elif kind == "table":
            headers = block.get("headers") or []
            rows = [[esc(c) for c in r] for r in (block.get("rows") or [])]
            data = ([[esc(h) for h in headers]] if headers else []) + rows
            if not data:
                continue
            table = Table(data, hAlign="LEFT")
            style = [
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#999999")),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
            if headers:
                style += [("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8E8E8")),
                          ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold")]
            table.setStyle(TableStyle(style))
            story += [table, Spacer(1, 10)]
            if block.get("caption"):
                story.append(Paragraph(esc(block["caption"]), caption_style))

        elif kind == "pagebreak":
            story.append(PageBreak())

    os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
    SimpleDocTemplate(
        path, pagesize=LETTER, title=title or "Assignment",
        leftMargin=0.9 * inch, rightMargin=0.9 * inch,
        topMargin=0.9 * inch, bottomMargin=0.9 * inch,
    ).build(story or [Spacer(1, 1)])
    return path
